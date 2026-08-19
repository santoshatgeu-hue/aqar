import tkinter as tk
from tkinter import ttk,filedialog,messagebox
from pathlib import Path
import pandas as pd, re, csv, hashlib, threading, queue, datetime, os, traceback

EXTS={".xlsx",".xls",".csv",".docx",".pdf",".txt"}
EXCLUDE_DIRS={"output",".git","__pycache__",".venv","venv"}
TEMP_PREFIXES=("~$",".~lock.")

def sha256(p):
    try:
        h=hashlib.sha256()
        with open(p,"rb") as f:
            for b in iter(lambda:f.read(1024*1024),b): h.update(b)
        return h.hexdigest()
    except: return ""

def load_csv(name):
    with open(Path(__file__).parent/name,encoding="utf-8-sig",newline="") as f:
        return list(csv.DictReader(f))

def metric_from_text(text,rules):
    s=str(text)
    # Exact known AQAR metric only; prevents false matches from arbitrary numbers.
    for r in rules:
        m=r["metric"].strip()
        if re.search(r"(?<!\d)"+re.escape(m)+r"(?!\d)",s,re.I): return m
    return ""

def excluded(p):
    return any(x in EXCLUDE_DIRS for x in p.parts) or p.name.startswith(TEMP_PREFIXES) or p.name=="AQAR_Master_Consolidation_v3_2.xlsx"

def dept_for(root,p):
    rel=p.relative_to(root)
    if len(rel.parts)==1: return root.name
    return rel.parts[0]

def criterion_for(p):
    for part in p.parts:
        m=re.match(r"(?i)criteria[-_ ]?(\d+)",part)
        if m: return "Criterion "+m.group(1)
    return ""

def evidence_type(p):
    return p.suffix.lower().replace(".","").upper() or "FILE"

class App:
    def __init__(self,root):
        self.root=root; self.q=queue.Queue(); self.rules=load_csv("metric_rules.csv")
        self.nt={r["metric"]:r for r in load_csv("non_template_requirements.csv")}
        self.selected=None; self.files=[]; self.scanning=False
        root.title("AQAR OneDrive Consolidator v3.2")
        root.geometry("980x680")
        ttk.Label(root,text="AQAR OneDrive Consolidator v3.2",font=("Arial",20,"bold")).pack(pady=(18,5))
        ttk.Label(root,text="Department data + Information Not in Data Template + supporting evidence").pack(pady=(0,15))
        f=ttk.LabelFrame(root,text="1. Select AQAR ROOT folder",padding=15); f.pack(fill="x",padx=25)
        self.path=tk.StringVar(value="No folder selected")
        ttk.Label(f,textvariable=self.path,wraplength=760).pack(side="left",fill="x",expand=True)
        ttk.Button(f,text="Select AQAR Folder",command=self.choose).pack(side="right")
        v=ttk.LabelFrame(root,text="2. Folder validation",padding=15); v.pack(fill="x",padx=25,pady=12)
        self.info=tk.StringVar(value="Select the university AQAR root containing department folders.")
        ttk.Label(v,textvariable=self.info).pack(anchor="w")
        self.bar=ttk.Progressbar(root,mode="determinate"); self.bar.pack(fill="x",padx=25)
        self.progress=tk.StringVar(value="Ready"); ttk.Label(root,textvariable=self.progress).pack(anchor="w",padx=25,pady=5)
        self.log=tk.Text(root,height=17,wrap="word"); self.log.pack(fill="both",expand=True,padx=25,pady=8)
        self.log.configure(state="disabled")
        b=ttk.Frame(root); b.pack(fill="x",padx=25,pady=15)
        self.startbtn=ttk.Button(b,text="Start Consolidation",command=self.start,state="disabled"); self.startbtn.pack(side="left")
        ttk.Button(b,text="Open Output Folder",command=self.open_output).pack(side="left",padx=10)
        ttk.Button(b,text="Exit",command=root.destroy).pack(side="right")
        root.after(100,self.poll)

    def write(self,s):
        self.log.configure(state="normal"); self.log.insert("end",s+"\n"); self.log.see("end"); self.log.configure(state="disabled")

    def choose(self):
        p=filedialog.askdirectory(title="Select university AQAR root folder")
        if not p:return
        self.selected=Path(p).resolve(); self.path.set(str(self.selected))
        self.files=[x for x in self.selected.rglob("*") if x.is_file() and x.suffix.lower() in EXTS and not excluded(x)]
        if not self.files:
            self.info.set("ERROR: no supported files found."); self.startbtn.configure(state="disabled")
            messagebox.showerror("No files","No supported files were found in this folder.")
            return
        deps=sorted(set(dept_for(self.selected,p) for p in self.files))
        self.info.set(f"{len(self.files):,} supported files found. Departments detected: {len(deps)}. Output/temporary files excluded.")
        self.startbtn.configure(state="normal")
        self.write("AQAR root: "+str(self.selected))
        self.write("Files found: "+str(len(self.files)))
        self.write("Departments detected: "+", ".join(deps[:30]))

    def start(self):
        if self.scanning:return
        self.scanning=True; self.startbtn.configure(state="disabled")
        threading.Thread(target=self.scan,daemon=True).start()

    def scan(self):
        out=self.selected/"output"; out.mkdir(exist_ok=True); logp=out/"aqar_scan.log"
        inv=[]; raw=[]; errors=[]
        with open(logp,"w",encoding="utf-8") as lf:
            def L(x): lf.write(x+"\n"); lf.flush(); self.q.put(("log",x))
            L("AQAR OneDrive Consolidator v3.2"); L("Started "+str(datetime.datetime.now()))
            L("Root "+str(self.selected))
            for n,p in enumerate(self.files,1):
                self.q.put(("progress",n,len(self.files),p.name))
                try:
                    h=sha256(p); d=dept_for(self.selected,p); c=criterion_for(p)
                    inv.append({"Department":d,"Criterion Folder":c,"Relative Path":str(p.relative_to(self.selected)),
                                "File":p.name,"Extension":p.suffix.lower(),"Size":p.stat().st_size,"SHA256":h})
                    if p.suffix.lower() in {".xlsx",".xls"}:
                        try:
                            xl=pd.ExcelFile(p)
                            for sh in xl.sheet_names:
                                try: df=pd.read_excel(p,sheet_name=sh,header=None)
                                except Exception as e:
                                    errors.append({"File":str(p),"Error":f"Sheet {sh}: {e}"}); continue
                                for i,row in df.iterrows():
                                    vals=[str(x) for x in row.tolist() if pd.notna(x)]
                                    if not vals: continue
                                    text=" | ".join(vals)
                                    m=metric_from_text(str(sh)+" "+text,self.rules)
                                    if m:
                                        rr=next(r for r in self.rules if r["metric"]==m)
                                        raw.append({"AQAR Metric":m,"Department":d,"Criterion":c,
                                          "Information Stream":"Data Template" if rr["non_template"]=="NO" else "Both / Review",
                                          "File":p.name,"Relative Path":str(p.relative_to(self.selected)),
                                          "Sheet":sh,"Row":i+1,"Content":" | ".join(vals[:30]),"File Hash":h,
                                          "Evidence Type":evidence_type(p)})
                        except Exception as e: errors.append({"File":str(p),"Error":str(e)})
                    elif p.suffix.lower()==".csv":
                        try:
                            df=pd.read_csv(p,header=None)
                            for i,row in df.iterrows():
                                vals=[str(x) for x in row.tolist() if pd.notna(x)]
                                m=metric_from_text(p.name+" "+" | ".join(vals),self.rules)
                                if m:
                                    raw.append({"AQAR Metric":m,"Department":d,"Criterion":c,
                                      "Information Stream":"Data/Supporting File","File":p.name,
                                      "Relative Path":str(p.relative_to(self.selected)),"Sheet":"","Row":i+1,
                                      "Content":" | ".join(vals[:30]),"File Hash":h,"Evidence Type":evidence_type(p)})
                        except Exception as e: errors.append({"File":str(p),"Error":str(e)})
                    else:
                        try:
                            text=""
                            if p.suffix.lower()==".txt": text=p.read_text(errors="ignore")
                            elif p.suffix.lower()==".pdf":
                                import pypdf
                                text="\n".join((x.extract_text() or "") for x in pypdf.PdfReader(str(p)).pages)
                            elif p.suffix.lower()==".docx":
                                from docx import Document
                                doc=Document(str(p)); text="\n".join(x.text for x in doc.paragraphs)
                                text+="\n"+"\n".join(" | ".join(cell.text for cell in row.cells) for table in doc.tables for row in table.rows)
                            m=metric_from_text(p.name+" "+text,self.rules)
                            if m:
                                raw.append({"AQAR Metric":m,"Department":d,"Criterion":c,
                                  "Information Stream":"Information Not in Data Template" if m in self.nt else "Supporting File",
                                  "File":p.name,"Relative Path":str(p.relative_to(self.selected)),
                                  "Sheet":"","Row":"","Content":text[:12000],"File Hash":h,"Evidence Type":evidence_type(p)})
                        except Exception as e: errors.append({"File":str(p),"Error":str(e)})
                except Exception as e: errors.append({"File":str(p),"Error":traceback.format_exc()})
            L(f"Completed files={len(inv)} records={len(raw)} errors={len(errors)}")
        rawdf=pd.DataFrame(raw)
        if rawdf.empty: rawdf=pd.DataFrame(columns=["AQAR Metric","Department","Criterion","Information Stream","File","Relative Path","Sheet","Row","Content","File Hash","Evidence Type"])
        invdf=pd.DataFrame(inv)
        # Supporting evidence inventory for the 12 metrics in supplied PDF.
        evidence_rows=[]
        for m,req in self.nt.items():
            g=rawdf[rawdf["AQAR Metric"]==m] if not rawdf.empty else rawdf
            depts=sorted(g["Department"].dropna().unique()) if not g.empty else []
            for d in sorted(set(dept_for(self.selected,p) for p in self.files)):
                gd=g[g["Department"]==d] if not g.empty else g
                evidence_rows.append({"Department":d,"AQAR Metric":m,"Required Information":req["required_information"],
                  "Submitted":"YES" if not gd.empty else "NO","Files":"; ".join(sorted(gd["Relative Path"].unique())) if not gd.empty else "",
                  "Evidence Types":", ".join(sorted(gd["Evidence Type"].unique())) if not gd.empty else "",
                  "IQAC Status":"VERIFY" if not gd.empty else "MISSING - ACTION REQUIRED"})
        evdf=pd.DataFrame(evidence_rows)
        present=set(rawdf["AQAR Metric"]); missing=[]
        for r in self.rules:
            missing.append({"AQAR Metric":r["metric"],"Title":r["title"],"Non-template":r["non_template"],
                            "Required Evidence":r["required_evidence"],"Found":r["metric"] in present})
        prop=[]
        for m,g in rawdf.groupby("AQAR Metric"):
            r=next((x for x in self.rules if x["metric"]==m),{})
            prop.append({"AQAR Metric":m,"Title":r.get("title",""),"Aggregation Rule":r.get("aggregation","REVIEW"),
                         "Departments":g["Department"].nunique(),"Records":len(g),
                         "Proposal":"PENDING IQAC APPROVAL","Status":"REVIEW REQUIRED"})
        proposed=pd.DataFrame(prop)
        dup=rawdf[rawdf.duplicated(["AQAR Metric","Content"],keep=False)] if not rawdf.empty else rawdf.copy()
        conflicts=[]
        for m,g in rawdf.groupby("AQAR Metric"):
            if len(g)>1:
                conflicts.append({"AQAR Metric":m,"Departments":g["Department"].nunique(),"Records":len(g),
                                  "Issue":"Multiple source records; apply metric-specific aggregation and deduplication","Status":"REVIEW REQUIRED"})
        x=out/"AQAR_Master_Consolidation_v3_2.xlsx"
        with pd.ExcelWriter(x,engine="openpyxl") as w:
            pd.DataFrame([{"AQAR Root":str(self.selected),"Files Scanned":len(invdf),"Records Extracted":len(rawdf),
                           "Metrics Found":rawdf["AQAR Metric"].nunique(),"Non-template Metrics":len(self.nt),
                           "Generated":str(datetime.datetime.now())}]).to_excel(w,index=False,sheet_name="ReadMe")
            invdf.to_excel(w,index=False,sheet_name="File Inventory")
            pd.DataFrame(self.rules).to_excel(w,index=False,sheet_name="Metric Rules")
            rawdf.to_excel(w,index=False,sheet_name="Raw Extract")
            proposed.to_excel(w,index=False,sheet_name="Proposed Consolidation")
            dup.to_excel(w,index=False,sheet_name="Duplicate Check")
            pd.DataFrame(conflicts).to_excel(w,index=False,sheet_name="Conflicts")
            pd.DataFrame(missing).to_excel(w,index=False,sheet_name="Missing Metrics")
            evdf.to_excel(w,index=False,sheet_name="Information Not in Data Template")
            rawdf[["AQAR Metric","Department","Criterion","Information Stream","File","Relative Path","Evidence Type","File Hash"]].drop_duplicates().to_excel(w,index=False,sheet_name="Evidence Mapping")
            pd.DataFrame(errors).to_excel(w,index=False,sheet_name="Errors")
        self.q.put(("done",str(x),len(invdf),len(rawdf),len(errors)))

    def poll(self):
        try:
            while True:
                x=self.q.get_nowait()
                if x[0]=="log": self.write(x[1])
                elif x[0]=="progress":
                    _,n,t,name=x; self.bar["maximum"]=t; self.bar["value"]=n; self.progress.set(f"Scanning {n:,}/{t:,}: {name}")
                elif x[0]=="done":
                    _,p,n,r,e=x; self.scanning=False; self.startbtn.configure(state="normal")
                    self.progress.set(f"Completed: {n:,} files, {r:,} records, {e:,} errors")
                    messagebox.showinfo("Completed",f"Files scanned: {n:,}\nRecords extracted: {r:,}\nErrors: {e:,}\n\nOutput:\n{p}")
        except queue.Empty: pass
        self.root.after(100,self.poll)

    def open_output(self):
        p=(self.selected/"output") if self.selected else Path(__file__).parent/"output"; p.mkdir(exist_ok=True)
        try:
            import subprocess; subprocess.Popen(["open" if os.uname().sysname=="Darwin" else "xdg-open",str(p)])
        except: messagebox.showinfo("Output",str(p))

root=tk.Tk(); App(root); root.mainloop()
